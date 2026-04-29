#!/usr/bin/env python3
"""
Универсальный конвертер MD → DOCX для договоров и юридических документов.
Основан на build_docx.py из client_board, адаптирован для контрактов.

Особенности:
- Multilevel нумерация Word (1., 1.1., 1.1.1. — автоматическая)
- Поддержка tracked changes: {++вставка++} и {--удаление--}
- Режим --accept: принять все правки (чистый текст)
- Режим --track: показать правки визуально (зелёный/красный)
- Сноски [^N] → footnotes
- Таблицы, нумерованные списки, заголовки
- Без логотипа/бренда (для внешних договоров)

Использование:
    python3 md_to_docx.py input.md                    # с tracked changes
    python3 md_to_docx.py input.md --accept            # принять все правки
    python3 md_to_docx.py input.md -o output.docx      # указать выходной файл
    python3 md_to_docx.py *.md                         # батч-конвертация
"""
import argparse
import os
import re
import sys

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# --- Colors ---
COLOR_HEADING = RGBColor(0x00, 0x00, 0x00)
COLOR_BODY = RGBColor(0x22, 0x22, 0x22)
COLOR_INSERT = RGBColor(0x00, 0x80, 0x00)   # green for insertions
COLOR_DELETE = RGBColor(0xFF, 0x00, 0x00)    # red for deletions
COLOR_COMMENT = RGBColor(0x00, 0x00, 0xCC)   # blue for comment refs

FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(11)
FONT_SIZE_SMALL = Pt(9)
FONT_SIZE_HEADING = Pt(13)
FONT_SIZE_SECTION = Pt(12)

# --- Page setup A4 ---
PAGE_WIDTH = Cm(21)
PAGE_HEIGHT = Cm(29.7)
MARGIN_LEFT = Cm(3)
MARGIN_RIGHT = Cm(1.5)
MARGIN_TOP = Cm(2)
MARGIN_BOTTOM = Cm(2)


# ---------------------------------------------------------------------------
# Tracked changes parsing
# ---------------------------------------------------------------------------

def parse_segments(text: str, accept: bool):
    """
    Parse text into segments: [(text, style), ...]
    style: 'normal', 'insert', 'delete', 'footnote'
    """
    segments = []
    combined = re.compile(
        r'(\{\+\+.*?\+\+\})'
        r'|(\{--.*?--\})'
        r'|(\[\^\d+\])',
        re.DOTALL
    )

    last = 0
    for m in combined.finditer(text):
        if m.start() > last:
            segments.append((text[last:m.start()], 'normal'))

        if m.group(1):  # insertion
            inner = m.group(1)[3:-3]
            if accept:
                segments.append((inner, 'normal'))
            else:
                segments.append((inner, 'insert'))
        elif m.group(2):  # deletion
            inner = m.group(2)[3:-3]
            if not accept:
                segments.append((inner, 'delete'))
        elif m.group(3):  # footnote ref
            if not accept:
                segments.append((m.group(3), 'footnote'))
            # in accept mode, footnote refs are omitted

        last = m.end()

    if last < len(text):
        segments.append((text[last:], 'normal'))

    return segments


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def is_table_row(line: str) -> bool:
    return line.strip().startswith("|") and line.strip().endswith("|")


def is_table_separator(line: str) -> bool:
    return bool(re.match(r"^\|[\s\-:|]+\|$", line.strip()))


def is_list_item(line: str) -> bool:
    return bool(re.match(r'^\s*[-*]\s', line))


def is_footnote_def(line: str) -> bool:
    return bool(re.match(r'^\[\^\d+\]:', line.strip()))


def _is_uppercase_text(text: str) -> bool:
    """Check if text (without numbering prefix) is all uppercase."""
    clean = re.sub(r'^\d+\.\s*', '', text.strip())
    # Strip tracked changes markers for checking
    clean = re.sub(r'\{\+\+.*?\+\+\}|\{--.*?--\}', '', clean)
    alpha = [c for c in clean if c.isalpha()]
    return len(alpha) > 2 and all(c.isupper() for c in alpha)


def _strip_number_prefix(text: str) -> str:
    """Strip leading 'N. ' from text."""
    return re.sub(r'^\d+\.\s*', '', text.strip())


def _get_leading_spaces(line: str) -> int:
    return len(line) - len(line.lstrip())


def _text_ends_with_colon(text: str) -> bool:
    """Check if text ends with ':' (ignoring tracked changes and whitespace)."""
    clean = re.sub(r'\{\+\+.*?\+\+\}|\{--.*?--\}|\[\^\d+\]', '', text)
    return clean.rstrip().endswith(':')


def _text_after_accept(text: str) -> str:
    """Return text as it would appear after accepting all changes."""
    # Remove deletions, keep insertion content
    result = re.sub(r'\{--.*?--\}', '', text)
    result = re.sub(r'\{\+\+(.*?)\+\+\}', r'\1', result)
    result = re.sub(r'\[\^\d+\]', '', result)
    return result.strip()


# ---------------------------------------------------------------------------
# Word multilevel numbering (XML)
# ---------------------------------------------------------------------------

def create_multilevel_numbering(doc):
    """Create a multilevel numbering definition for legal contracts.

    Levels:
        0: "1."      — chapters (bold, no indent)
        1: "1.1."    — clauses (no indent)
        2: "1.1.1."  — sub-clauses (~1cm indent)
        3: "1.1.1.1." — deeper (~2cm indent)
        4: "1.1.1.1.1." — deepest (~3cm indent)

    Returns the numId to use in paragraphs.
    """
    numbering_part = doc.part.part_related_by(RT.NUMBERING)
    numbering_elm = numbering_part.element

    ABSTRACT_NUM_ID = 10  # above default 0-8

    abstract = OxmlElement('w:abstractNum')
    abstract.set(qn('w:abstractNumId'), str(ABSTRACT_NUM_ID))

    nsid = OxmlElement('w:nsid')
    nsid.set(qn('w:val'), '1A2B3C4D')
    abstract.append(nsid)

    mlt = OxmlElement('w:multiLevelType')
    mlt.set(qn('w:val'), 'multilevel')
    abstract.append(mlt)

    tmpl = OxmlElement('w:tmpl')
    tmpl.set(qn('w:val'), '5E6F7A8B')
    abstract.append(tmpl)

    # Level configs: (ilvl, lvlText, left_twips, hanging_twips, bold)
    levels = [
        (0, '%1.', 425, 425, True),          # "1." chapters
        (1, '%1.%2.', 0, 0, False),          # "1.1." clauses — no indent
        (2, '%1.%2.%3.', 567, 567, False),   # "1.1.1." sub-clauses
        (3, '%1.%2.%3.%4.', 1134, 567, False),
        (4, '%1.%2.%3.%4.%5.', 1701, 567, False),
    ]

    for ilvl, lvl_text, left, hanging, bold in levels:
        lvl = OxmlElement('w:lvl')
        lvl.set(qn('w:ilvl'), str(ilvl))

        start = OxmlElement('w:start')
        start.set(qn('w:val'), '1')
        lvl.append(start)

        num_fmt = OxmlElement('w:numFmt')
        num_fmt.set(qn('w:val'), 'decimal')
        lvl.append(num_fmt)

        lvl_text_el = OxmlElement('w:lvlText')
        lvl_text_el.set(qn('w:val'), lvl_text)
        lvl.append(lvl_text_el)

        lvl_jc = OxmlElement('w:lvlJc')
        lvl_jc.set(qn('w:val'), 'left')
        lvl.append(lvl_jc)

        pPr = OxmlElement('w:pPr')
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), str(left))
        ind.set(qn('w:hanging'), str(hanging))
        pPr.append(ind)
        lvl.append(pPr)

        if bold:
            rPr = OxmlElement('w:rPr')
            b_el = OxmlElement('w:b')
            bCs = OxmlElement('w:bCs')
            rPr.append(b_el)
            rPr.append(bCs)
            lvl.append(rPr)

        abstract.append(lvl)

    # Insert abstractNum before first w:num (schema requirement)
    first_num = numbering_elm.find(qn('w:num'))
    if first_num is not None:
        numbering_elm.insert(list(numbering_elm).index(first_num), abstract)
    else:
        numbering_elm.append(abstract)

    # Add w:num referencing it
    num = OxmlElement('w:num')
    num.set(qn('w:numId'), '10')
    abstract_ref = OxmlElement('w:abstractNumId')
    abstract_ref.set(qn('w:val'), str(ABSTRACT_NUM_ID))
    num.append(abstract_ref)
    numbering_elm.append(num)

    return 10  # numId


def set_paragraph_numbering(para, num_id: int, ilvl: int):
    """Set Word numbering on a paragraph."""
    pPr = para._p.get_or_add_pPr()
    # Remove any existing numPr
    existing = pPr.find(qn('w:numPr'))
    if existing is not None:
        pPr.remove(existing)

    numPr = OxmlElement('w:numPr')
    ilvl_elm = OxmlElement('w:ilvl')
    ilvl_elm.set(qn('w:val'), str(ilvl))
    numId_elm = OxmlElement('w:numId')
    numId_elm.set(qn('w:val'), str(num_id))
    numPr.append(ilvl_elm)
    numPr.append(numId_elm)
    pPr.append(numPr)


# ---------------------------------------------------------------------------
# NumberingDetector — state machine
# ---------------------------------------------------------------------------

class NumberingDetector:
    """Detects the ilvl for each line based on MD structure.

    Returns (line_type, ilvl, clean_text) where:
        line_type: 'chapter', 'clause', 'subclause', 'bullet', 'text', 'empty'
        ilvl: 0, 1, 2 (or -1 for non-numbered)
        clean_text: text with number prefix stripped
    """

    def __init__(self):
        self.parent_ends_with_colon = False
        self.in_subclause_block = False
        self.last_clause_number = 0   # tracks clause-level numbering
        self.next_chapter_number = 1  # tracks expected next chapter number

    def _enter_chapter(self, text):
        """Common logic when entering a chapter."""
        self.parent_ends_with_colon = False
        self.in_subclause_block = False
        self.last_clause_number = 0
        self.next_chapter_number += 1
        return ('chapter', 0, text)

    def classify(self, line: str):
        stripped = line.strip()

        if not stripped:
            return ('empty', -1, '')

        # Bullet items (- text)
        if re.match(r'^\s*[-*]\s', line):
            text = re.sub(r'^\s*[-*]\s', '', line).strip()
            return ('bullet', -1, text)

        # Clean "## " artifacts from numbered lines (e.g. "25. ## В случае...")
        stripped = re.sub(r'^(\d+\.)\s*##\s*', r'\1 ', stripped)

        # Check for numbered item: "N. text"
        num_match = re.match(r'^(\d+)\.\s+(.+)$', stripped)
        if not num_match:
            # Plain text paragraph
            if len(stripped) > 30:
                self.in_subclause_block = False
            self.parent_ends_with_colon = False
            return ('text', -1, stripped)

        number = int(num_match.group(1))
        text = num_match.group(2)

        # Check if UPPERCASE heading → chapter (ilvl=0)
        if _is_uppercase_text(stripped):
            return self._enter_chapter(text)

        # Check if chapter by number match: N == expected next chapter
        # and text looks like a title (short, no indent, not a regular clause)
        if (number == self.next_chapter_number
                and _get_leading_spaces(line) == 0
                and len(text) < 80
                and not _text_ends_with_colon(text)):
            return self._enter_chapter(text)

        # Key heuristic: use last_clause_number to distinguish
        # clause continuation from sub-clauses.
        #
        # After a clause ending with ":", if the next number is NOT
        # last_clause_number + 1, it's a sub-clause.
        # When in a sub-clause block, number == last_clause_number + 1
        # means we're back to clause level.

        next_clause = self.last_clause_number + 1

        # Start sub-clause block?
        if self.parent_ends_with_colon and number != next_clause:
            self.in_subclause_block = True
            self.parent_ends_with_colon = _text_ends_with_colon(text)
            return ('subclause', 2, text)

        # Continue sub-clause block?
        if self.in_subclause_block:
            if number == next_clause:
                # Break out — next clause in sequence
                self.in_subclause_block = False
                self.last_clause_number = number
                self.parent_ends_with_colon = _text_ends_with_colon(text)
                return ('clause', 1, text)
            # Stay in sub-clause block
            self.parent_ends_with_colon = _text_ends_with_colon(text)
            return ('subclause', 2, text)

        # Regular clause (ilvl=1)
        self.last_clause_number = number
        self.parent_ends_with_colon = _text_ends_with_colon(text)
        return ('clause', 1, text)


# ---------------------------------------------------------------------------
# Paragraph building
# ---------------------------------------------------------------------------

def add_run_with_style(para, text: str, style: str, font_name: str,
                       font_size, accept: bool):
    """Add a styled run to paragraph."""
    if not text:
        return

    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = font_size

    if style == 'normal':
        run.font.color.rgb = COLOR_BODY
    elif style == 'insert':
        run.font.color.rgb = COLOR_INSERT
        run.font.underline = True
    elif style == 'delete':
        run.font.color.rgb = COLOR_DELETE
        run.font.strike = True
    elif style == 'footnote':
        run.font.color.rgb = COLOR_COMMENT
        run.font.size = FONT_SIZE_SMALL
        run.font.superscript = True


def add_paragraph(doc, text: str, accept: bool, bold: bool = False,
                  font_size=None, color: RGBColor = None,
                  alignment=None, indent_cm: float = 0):
    """Add a paragraph with tracked changes support."""
    para = doc.add_paragraph()
    if alignment:
        para.alignment = alignment

    if indent_cm > 0:
        para.paragraph_format.left_indent = Cm(indent_cm)

    if not text:
        return para

    fs = font_size or FONT_SIZE
    segments = parse_segments(text, accept)

    for seg_text, seg_style in segments:
        if bold and seg_style == 'normal':
            run = para.add_run(seg_text)
            run.font.name = FONT_NAME
            run.font.size = fs
            run.font.bold = True
            run.font.color.rgb = color or COLOR_HEADING
        else:
            add_run_with_style(para, seg_text, seg_style, FONT_NAME, fs, accept)

    return para


def _set_table_borders_none(table):
    """Remove all borders from a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'none')
        el.set(qn('w:sz'), '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'auto')
        borders.append(el)
    # Remove existing borders if any
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)


def add_table(doc, rows_data, accept: bool):
    """Add a table from parsed rows.

    2-column tables get special treatment: no borders,
    right column right-aligned (contract header/footer style).
    """
    if not rows_data:
        return

    num_cols = max(len(r) for r in rows_data)
    for r in rows_data:
        while len(r) < num_cols:
            r.append('')

    table = doc.add_table(rows=len(rows_data), cols=num_cols)

    # 2-column tables: no borders
    is_two_col = (num_cols == 2)
    if is_two_col:
        _set_table_borders_none(table)
    else:
        table.style = 'Table Grid'

    # Short 2-col table (1 row, short cells) = header table (city/date)
    # → right-align last column
    is_header_table = (is_two_col and len(rows_data) == 1
                       and all(len(c) < 80 for c in rows_data[0]))

    for i, row_data in enumerate(rows_data):
        for j, cell_text in enumerate(row_data):
            cell = table.cell(i, j)
            cell.text = ""
            # Split on <br> for multi-line cells
            cell_lines = cell_text.strip().split('<br>')
            for k, cl in enumerate(cell_lines):
                if k == 0:
                    para = cell.paragraphs[0]
                else:
                    para = cell.add_paragraph()
                # Right-align last column only in header table
                if is_header_table and j == num_cols - 1:
                    from docx.enum.text import WD_ALIGN_PARAGRAPH
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                segments = parse_segments(cl.strip(), accept)
                for seg_text, seg_style in segments:
                    add_run_with_style(para, seg_text, seg_style, FONT_NAME,
                                       FONT_SIZE_SMALL, accept)

    return table


def parse_table_rows(lines):
    """Parse table lines into list of lists."""
    rows = []
    for line in lines:
        if is_table_separator(line):
            continue
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    return rows


# ---------------------------------------------------------------------------
# Main build
# ---------------------------------------------------------------------------

def build_docx(md_path: str, docx_path: str, accept: bool = False):
    """Convert MD to DOCX with multilevel numbering."""
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.read().split("\n")

    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    section.left_margin = MARGIN_LEFT
    section.right_margin = MARGIN_RIGHT
    section.top_margin = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM

    # --- Default style ---
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = FONT_SIZE
    style.font.color.rgb = COLOR_BODY
    pf = style.paragraph_format
    pf.space_after = Pt(2)
    pf.space_before = Pt(2)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # --- Create multilevel numbering ---
    num_id = create_multilevel_numbering(doc)
    detector = NumberingDetector()

    # --- First non-empty line = document title (bold, centered) ---
    first_line_idx = 0
    while first_line_idx < len(lines) and not lines[first_line_idx].strip():
        first_line_idx += 1
    if first_line_idx < len(lines):
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        add_paragraph(doc, lines[first_line_idx].strip(), accept, bold=True,
                      font_size=FONT_SIZE_HEADING, color=COLOR_HEADING,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)
        first_line_idx += 1
    else:
        first_line_idx = 0

    # --- Process lines ---
    i = first_line_idx
    while i < len(lines):
        line = lines[i].rstrip()

        # Skip the legend block at the end
        if line.strip() == '**Обозначения tracked changes:**':
            break
        if line.strip() == '**Комментарии рецензента:**':
            if not accept:
                add_paragraph(doc, 'Комментарии рецензента:', accept,
                              bold=True, font_size=FONT_SIZE_SECTION)
            i += 1
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Horizontal rule
        if line.strip() == '---':
            i += 1
            continue

        # Footnote definitions
        if is_footnote_def(line):
            if not accept:
                fn_text = line.strip()
                add_paragraph(doc, fn_text, accept, font_size=FONT_SIZE_SMALL,
                              indent_cm=1.0)
            i += 1
            continue

        # Table block
        if is_table_row(line):
            table_lines = []
            while i < len(lines) and is_table_row(lines[i].rstrip()):
                table_lines.append(lines[i].rstrip())
                i += 1
            rows_data = parse_table_rows(table_lines)
            add_table(doc, rows_data, accept)
            continue

        # Classify the line
        line_type, ilvl, clean_text = detector.classify(line)

        if line_type == 'empty':
            i += 1
            continue

        # Skip lines that become empty after accepting changes
        if accept and line_type in ('clause', 'subclause', 'text') \
                and not _text_after_accept(clean_text):
            i += 1
            continue

        if line_type == 'chapter':
            # Chapter heading — ilvl=0, bold, section font
            para = add_paragraph(doc, clean_text, accept, bold=True,
                                 font_size=FONT_SIZE_SECTION, color=COLOR_HEADING)
            set_paragraph_numbering(para, num_id, 0)
            # Add some spacing before chapter
            para.paragraph_format.space_before = Pt(12)
            i += 1
            continue

        if line_type == 'clause':
            # Clause — ilvl=1, normal text
            para = add_paragraph(doc, clean_text, accept)
            set_paragraph_numbering(para, num_id, 1)
            i += 1
            continue

        if line_type == 'subclause':
            # Sub-clause — ilvl=2
            para = add_paragraph(doc, clean_text, accept)
            set_paragraph_numbering(para, num_id, 2)
            i += 1
            continue

        if line_type == 'bullet':
            para = add_paragraph(doc, clean_text, accept, indent_cm=1.5)
            # Add bullet run at the start
            first_run = para.runs[0] if para.runs else None
            if first_run:
                first_run.text = '– ' + first_run.text
            i += 1
            continue

        # Plain text paragraph
        para = add_paragraph(doc, clean_text, accept)
        i += 1

    # --- Save ---
    doc.save(docx_path)
    print(f"Generated: {docx_path}")
    return docx_path


def main():
    parser = argparse.ArgumentParser(
        description='Convert Markdown to DOCX for contracts/legal documents'
    )
    parser.add_argument('files', nargs='+', help='MD file(s) to convert')
    parser.add_argument('-o', '--output', help='Output file path (single file only)')
    parser.add_argument('--accept', action='store_true',
                        help='Accept all tracked changes (clean document)')
    parser.add_argument('--track', action='store_true',
                        help='Show tracked changes visually (default)')
    args = parser.parse_args()

    accept = args.accept

    if args.output and len(args.files) > 1:
        print("Error: -o/--output can only be used with a single file", file=sys.stderr)
        sys.exit(1)

    for md_path in args.files:
        p = md_path
        if not os.path.exists(p):
            print(f"Error: {p} not found", file=sys.stderr)
            continue

        if args.output:
            out_path = args.output
        else:
            base = os.path.splitext(p)[0]
            suffix = '_clean' if accept else ''
            out_path = f"{base}{suffix}.docx"

        build_docx(p, out_path, accept=accept)


if __name__ == '__main__':
    main()
